#!/usr/bin/env python3
"""
Convertir Manual de Usuario Markdown a Word (.docx)
Preserva estructura, títulos, listas, tablas y formato
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def markdown_to_docx(markdown_file, output_file):
    """Convierte un archivo Markdown a documento Word con formato"""

    # Leer el archivo markdown
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Crear documento Word
    doc = Document()

    # Configurar estilos del documento
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Procesar línea por línea
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Saltar líneas vacías (excepto entre párrafos importantes)
        if not line.strip():
            i += 1
            continue

        # === TÍTULOS ===
        if line.startswith('# '):
            # Título nivel 1
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1

        elif line.startswith('## '):
            # Título nivel 2
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            i += 1

        elif line.startswith('### '):
            # Título nivel 3
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            i += 1

        elif line.startswith('#### '):
            # Título nivel 4
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
            i += 1

        # === LÍNEAS DE SEPARACIÓN ===
        elif line.strip() == '---':
            doc.add_paragraph()
            i += 1

        # === TABLAS ===
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Detectar tabla markdown
            table_lines = []
            j = i

            # Recopilar líneas de tabla
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1

            if len(table_lines) > 1:
                # Procesar tabla
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

                # Crear tabla en Word
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'

                # Agregar encabezados
                header_cells = table.rows[0].cells
                for idx, header in enumerate(headers):
                    header_cells[idx].text = header
                    # Formato de encabezado
                    for paragraph in header_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # Agregar datos (saltar línea de separador)
                start_data = 2 if len(table_lines) > 2 else 1
                for row_idx in range(start_data, len(table_lines)):
                    cells_text = [cell.strip() for cell in table_lines[row_idx].split('|')[1:-1]]
                    row_cells = table.add_row().cells
                    for col_idx, cell_text in enumerate(cells_text):
                        row_cells[col_idx].text = cell_text

                i = j
            else:
                i += 1

        # === LISTAS CON VIÑETAS ===
        elif line.strip().startswith('- '):
            text = line.strip()[2:].strip()
            # Procesar formato de texto
            text = procesar_formato_texto(text)
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1

        # === LISTAS NUMERADAS ===
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = procesar_formato_texto(text)
            p = doc.add_paragraph(text, style='List Number')
            i += 1

        # === PÁRRAFOS NORMALES ===
        else:
            # Procesar formato de texto (negritas, cursivas, etc.)
            text = procesar_formato_texto(line.strip())

            if text:
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            i += 1

    # Guardar documento
    doc.save(output_file)
    print(f"✅ Documento convertido exitosamente: {output_file}")

def procesar_formato_texto(text):
    """Procesa formato markdown en el texto (negritas, cursivas, etc.)"""

    # Preservar el texto para procesamiento
    # Nota: python-docx no soporta directamente markdown,
    # así que solo preservamos el texto sin el markdown

    # Remover enlaces markdown [texto](url)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

    # Remover formato markdown
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Negritas
    text = re.sub(r'\*([^*]+)\*', r'\1', text)      # Cursivas
    text = re.sub(r'`([^`]+)`', r'\1', text)        # Código

    return text

if __name__ == "__main__":
    source = Path("c:\\Personal\\CLIP_Comparador_V2\\docs\\MANUAL_USUARIO_TIENDA.md")
    output = Path("c:\\Personal\\CLIP_Comparador_V2\\docs\\MANUAL_USUARIO_TIENDA.docx")

    if source.exists():
        print(f"📄 Convirtiendo: {source}")
        markdown_to_docx(str(source), str(output))
        print(f"📍 Ubicación: {output}")
    else:
        print(f"❌ Archivo no encontrado: {source}")
